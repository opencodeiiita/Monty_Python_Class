import docx
import os
import shutil

print("")
date = str(input("Please enter today's date: "))
name = str(input("Please enter your full name: ")).upper()
roll_number = str(input("Please enter your roll number: "))
year = str(input("Please enter year: "))
leave_days = str(input("Please enter number of requested leave days: "))
start_day = str(input("Please enter start day of the leave: "))
end_day = str(input("Please enter end day of the leave: "))
print("")

# Starting connection with the template
orignal_file = 'Leave Letter Format.docx'
original_path = os.path.join(os.getcwd(), orignal_file)
doc = docx.Document(original_path)

# Inserting date, name, roll number and year
date_par = doc.paragraphs[2].add_run(' {0}'.format(date))
name_par = doc.paragraphs[4].add_run(' {0}'.format(name))
roll_number_par = doc.paragraphs[5].add_run(' {0}'.format(roll_number))
year_par = doc.paragraphs[6].add_run(' {0}'.format(year))

# Editing main part of the template including leave days, start day and end day
main_text = doc.paragraphs[16].text
main_text = main_text.replace('____', leave_days, 1)
main_text = main_text.replace('____________', start_day, 1)
main_text = main_text.replace('___________', end_day, 1)
doc.paragraphs[16].text = main_text

# Saving a new file
new_file_name = 'Leave Letter {0}.docx'.format(name)
path_new_file = os.path.join(os.getcwd(), new_file_name)
shutil.copy2(original_path, path_new_file) # complete target filename given
doc.save(new_file_name)

# Checking output file (new file)
doc2 = docx.Document(new_file_name)
num_par = doc.paragraphs
for i in range(len(num_par)):
    print(doc2.paragraphs[i].text)